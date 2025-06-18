import os
from google.cloud import storage
from docx import Document as DocxDocument
from PyPDF2 import PdfReader
from langchain_core.documents import Document
from langchain.chains import LLMChain
from langchain.prompts import PromptTemplate
from langchain_google_genai import ChatGoogleGenerativeAI
import tempfile
from google.cloud import storage
import uuid
from pathlib import Path
from process_design import *

# IMPORTANT: Replace these with your actual GCP Project ID, Region, and GCS Bucket.
PROJECT_ID = "hacker2025-team-97-dev" 
REGION = "us-central1"              
GCS_BUCKET_NAME = "hacker2025-team-97-dev.appspot.com" 

storage_client = storage.Client(project=PROJECT_ID)
SOURCE_REQUIREMENT_FOLDER = "source documents/"
# ----- Set up Google Gemini -----
genai = ChatGoogleGenerativeAI(
    model="gemini-2.0-flash",
    temperature=0,
    max_tokens=None,
    timeout=None,
    max_retries=2
)

# ----- File download from GCS -----
def download_file_from_gcs(bucket_name, blob_name, local_path):
    bucket = storage_client.bucket(bucket_name)
    blob = bucket.blob(blob_name)
    blob.download_to_filename(local_path)
    print(f"Downloaded {blob_name} to {local_path}")

# ----- Load and chunk DOCX -----
def load_and_chunk_docx_with_metadata(file_path, file_name):
    doc = DocxDocument(file_path)
    docs = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            docs.append(Document(
                page_content=text,
                metadata={"source_file": file_name, "chunk_number": i+1}
            ))
    return docs

# ----- Load and chunk PDF -----
def load_and_chunk_pdf_with_metadata(file_path, file_name):
    reader = PdfReader(file_path)
    docs = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text:
            docs.append(Document(
                page_content=text.strip(),
                metadata={"source_file": file_name, "page_number": i+1}
            ))
    return docs

# ----- Main pipeline -----
def process_file(bucket_name, blob_name, local_path):
    # Step 1: Download
    download_file_from_gcs(bucket_name, blob_name, local_path)
    file_ext = os.path.splitext(local_path)[1].lower()

    # Step 2: Chunk
    if file_ext == ".docx":
        docs = load_and_chunk_docx_with_metadata(local_path, blob_name)
    elif file_ext == ".pdf":
        docs = load_and_chunk_pdf_with_metadata(local_path, blob_name)
    else:
        raise ValueError("Unsupported file type. Only .docx and .pdf are supported.")

    # Step 3: Extract Requirements
    prompt = PromptTemplate(
        input_variables=["text"],
        template="""
        Extract the requirements info from the following text
        {text}
        Extract all clear, meaningful requirements and return them as an array of strings under the "requirements" key.

        Format the response in **strict JSON** like this:
        {{
        "requirements": [
            "First requirement goes here.",
            "Second requirement goes here.",
            "... more if present"
        ]
        }}

        Do not include markdown formatting like triple backticks. Only return valid JSON string.
        """
    )

    # Initialize your chain with the prompt and the Google GenAI chat model
    extraction_chain = prompt | genai
    all_requirements = []
    for doc in docs:
        output_text = extraction_chain.invoke(doc.page_content)
        
        # Parse the output JSON text safely
        import json
        try:
            # Parse the JSON string into a Python dictionary
            extracted_items = json.loads(output_text.content[8:-4])
        except json.JSONDecodeError as e:
            print("JSON decode failed:", e)
            extracted_items = {}

        # Add metadata info to each extracted item
        if "requirements" in extracted_items:
            for item in extracted_items["requirements"]:
                requirement = {
                    "requirement": item,
                    "source_file": doc.metadata.get("source_file"),
                    "page_number": doc.metadata.get("page_number") or doc.metadata.get("chunk_number")
                }
                all_requirements.append(requirement)


    # Step 4: Assign IDs
    for idx, req in enumerate(all_requirements):
        req["requirement_id"] = f"REQ-{idx+len(extracted_requirements)+1:04d}"
    return all_requirements

# ----- Entry point -----
if __name__ == "__main__":
    print(f"Starting Requirement Validation Agent for Project: {PROJECT_ID}, Region: {REGION}")
    print(f"GCS Bucket: {GCS_BUCKET_NAME}")
    extracted_requirements = []
    # Create a temporary directory for downloaded files
    with tempfile.TemporaryDirectory() as tmpdir:
        tmp_dir_path = tmpdir+"/downloads/requirements"
        os.makedirs(tmp_dir_path, exist_ok=True)

        print("\n--- Step 1: Document Ingestion and Preprocessing ---")
        
        print("\n--- Step 1A: Fetching all requirement files ---")
        # Downloading source documents from GCS¸
        source_docs = []
        blobs = storage_client.list_blobs(GCS_BUCKET_NAME, prefix=SOURCE_REQUIREMENT_FOLDER)

        for blob in blobs:
            if blob.name.endswith("/") or not any(blob.name.endswith(ext) for ext in [".pdf", ".docx", ".txt"]):
                continue  # Skip folders or unsupported files
            print(f"Processing {blob.name}")
            unique_suffix = uuid.uuid4().hex[:8]
            original_name = Path(blob.name).name
            unique_name = f"{unique_suffix}_{original_name}"
            local_filename = tmp_dir_path+"/"+unique_name
            extracted_requirements.extend(process_file(GCS_BUCKET_NAME, blob.name, local_filename))
        print("Total requirements extracted: ", len(extracted_requirements))
        process_design([item["requirement"] for item in extracted_requirements])


            


            